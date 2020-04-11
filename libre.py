import os
import subprocess
import glob
import re
import json
import sys
from docx import Document

Path = '/Users/aman/Desktop/pdfutility/resumes'

list_of_files = glob.glob(Path + '/*')
inputFilename = max(list_of_files, key=os.path.getctime)
fileWithoutExt = os.path.splitext(inputFilename)[0]
extension = os.path.splitext(inputFilename)[1]
# print(fileWithoutExt)

def convertToDoc():
	try:
		subprocess.call('soffice --headless --infilter=writer_pdf_import --convert-to docx --outdir /Users/aman/Desktop/pdfutility/resumes "{}"'
                    .format(inputFilename), shell=True)
	except:
		print("ERROR CONVERTING")


def getInfo(bolds,emails,phones,headings):
	f = open(fileWithoutExt+'.docx', 'rb')
	document = Document(f)
	all_paras = document.paragraphs

	bad_chars = [';', ':', '!', '*', '(',')', ',','.']
	for para in document.paragraphs:
		text = para.text
		email_list = re.findall(r"[a-z0-9\.\-+_]+@[a-z0-9\.\-+_]+\.[a-z]+",text)
		phone_list=re.findall(r'[\+\(]?[0-9][0-9 .\-\(\)]{8,}[0-9]',text)

		for email in email_list:
			email = re.sub(' +', ' ', email)
			for i in bad_chars : 
				email = email.replace(i, '')
				emails.append(email.strip())
		
		for phone in phone_list:
			phone = re.sub(' +', ' ', phone)
			for i in bad_chars :
				phone = phone.replace(i, '')
				phones.append(phone.strip())
		
		for run in para.runs:
			if run.bold:
				run.text = re.sub(' +', ' ',run.text)
			    for i in bad_chars :
				    run.text = run.text.replace(i, '')
				    bolds.append(run.text.strip())
		
		if para.style.name == 'Heading 1':
			para.text = re.sub(' +', ' ',para.text)
			for i in bad_chars :
				para.text = para.text.replace(i, '')
				headings.append(para.text)
		

		emails = list(filter(None, list(map(lambda s: s.strip(), emails))))
		phones = list(filter(None,list(map(lambda s: s.strip(), phones))))
		bolds = list(filter(None,list(map(lambda s: s.strip(), bolds))))
		headings = list(filter(None,list(map(lambda s: s.strip(), headings))))
		
		

		emails = [x.strip('"') for x in emails]
		phones = [i.strip('"') for i in phones]
		bolds = [i.strip('"') for i in bolds]
		headings = [i.strip('"') for i in headings]



if extension == '.pdf':
	convertToDoc()

bolds=[]
emails=[]
phones=[]
headings=[]

getInfo(bolds,emails,phones,headings)

style_Dict={'emails':emails,'phone_numbers':phones,'bold_phrases':bolds,'headings':headings}
 

r = json.dumps(style_Dict)
loaded_r = json.loads(r)

# print(json.dumps(loaded_r,sort_keys=False))

print(loaded_r)
sys.stdout.flush()